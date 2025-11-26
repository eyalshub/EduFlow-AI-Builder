import os
import re
import json
import ftfy
from datetime import datetime
from typing import Optional
from docx import Document
from docx.table import Table
from docx.oxml.ns import qn
from docx.shared import Inches

class DocxExtractor:
    def __init__(self, path: str):
        self.path = path
        self.filename = os.path.basename(path)
        self.basename = os.path.splitext(self.filename)[0]
        self.text = ""
        self.metadata = {}
        self.document = Document(path)

    def is_hebrew(self, text: str, threshold: float = 0.2) -> bool:
        hebrew_chars = re.findall(r'[\u0590-\u05FF]', text)
        return (len(hebrew_chars) / len(text)) >= threshold if text else False

    def fix_encoding(self, text: str) -> str:
        if not text:
            return ""
        if "�" in text or "\u202e" in text or sum(1 for c in text if ord(c) > 127) / len(text) > 0.3:
            return ftfy.fix_text(text).strip()
        return text.strip()

    def reverse_hebrew_sentence_words(self, text: str) -> str:
        def is_hebrew_line(line):
            return re.search(r'[\u0590-\u05FF]', line) is not None

        fixed_lines = []
        for line in text.splitlines():
            line = line.strip()
            if is_hebrew_line(line):
                trailing_punct = ""
                if line and line[-1] in ".!?":
                    trailing_punct = line[-1]
                    line = line[:-1]

                words = line.split()
                reversed_line = " ".join(reversed(words)) + trailing_punct
                fixed_lines.append(reversed_line)
            else:
                fixed_lines.append(line)
        return "\n".join(fixed_lines)


    def extract_text(self) -> str:
        try:
            paragraphs = [p.text.strip() for p in self.document.paragraphs if p.text.strip()]
            raw_text = "\n\n".join(paragraphs)
            clean_text = self.fix_encoding(raw_text)
            if self.is_hebrew(clean_text):
                clean_text = self.reverse_hebrew_sentence_words(clean_text)
            self.text = clean_text.strip()
            return self.text
        except Exception as e:
            print(f"[ERROR] Failed to extract DOCX text: {e}")
            self.text = ""
            return ""

    def extract_tables(self, output_path: str):
        tables_dir = os.path.join(output_path, "tables")
        os.makedirs(tables_dir, exist_ok=True)
        for i, table in enumerate(self.document.tables, start=1):
            table_text = ""
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text += " | ".join(row_text) + "\n"
            with open(os.path.join(tables_dir, f"table_{i}.txt"), "w", encoding="utf-8") as f:
                f.write(table_text)

    def extract_images(self, output_path: str):
        images_dir = os.path.join(output_path, "images")
        os.makedirs(images_dir, exist_ok=True)
        rels = self.document.part._rels
        img_count = 0
        for rel in rels:
            rel_obj = rels[rel]
            if "image" in rel_obj.reltype:
                img_count += 1
                img_data = rel_obj.target_part.blob
                ext = os.path.splitext(rel_obj.target_ref)[-1]
                with open(os.path.join(images_dir, f"image_{img_count}{ext}"), "wb") as f:
                    f.write(img_data)

    def save(self, output_dir="extracted_data"):
        output_path = os.path.join(output_dir, self.basename)
        os.makedirs(output_path, exist_ok=True)

        # Save content
        with open(os.path.join(output_path, "content.txt"), "w", encoding="utf-8") as f:
            f.write(self.text if self.text else "[EMPTY]")

        # Save metadata
        self.metadata = {
            "filename": self.filename,
            "word_count": len(self.text.split()) if self.text else 0,
            "timestamp": datetime.now().isoformat(),
            "path": self.path,
            "status": "success" if self.text else "empty",
            "language": "hebrew" if self.is_hebrew(self.text) else "unknown"
        }

        with open(os.path.join(output_path, "meta.json"), "w", encoding="utf-8") as f:
            json.dump(self.metadata, f, indent=2, ensure_ascii=False)

        print(f"[INFO] Saved text and metadata to: {output_path}")

        # Save images and tables
        self.extract_images(output_path)
        self.extract_tables(output_path)
        print("[INFO] Images and tables saved.")


if __name__ == "__main__":
    path = input("Enter DOCX file path: ").strip()
    if os.path.isfile(path):
        extractor = DocxExtractor(path)
        extractor.extract_text()
        extractor.save()
    else:
        print("[ERROR] File not found.")
